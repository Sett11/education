#!/usr/bin/env python3
"""
Простой процессор Word документов на python-docx
Тестирует сохранение форматирования при изменении содержимого
"""

import os
import traceback
from datetime import datetime
from typing import Dict, List
from dotenv import load_dotenv

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from mylogger import Logger
from llm import OpenRouterClient

# Загружаем переменные окружения из .env файла
load_dotenv()

class WordProcessor:
    """Простой процессор Word документов с сохранением форматирования"""
    
    def __init__(self):
        self.document = None
        self.logger = Logger('WordProcessor', 'logs/word_processor.log')
        self.logger.info("WordProcessor инициализирован")
    
    def load_document(self, filepath: str) -> 'WordProcessor':
        """Загружает документ из файла"""
        print(f"📖 Загружаем документ: {filepath}")
        self.logger.info(f"Загружаем документ: {filepath}")
        try:
            self.document = Document(filepath)
            print("✅ Документ загружен успешно")
            self.logger.info(f"Документ загружен успешно: {filepath}")
        except Exception as e:
            self.logger.error(f"Ошибка при загрузке документа {filepath}: {e}")
            raise
        return self
    
    def process_document(self, input_filepath: str, prompt: str, llm_client=None) -> str:
        """
        Обрабатывает документ с помощью LLM: читает, отправляет в модель и сохраняет результат
        
        Args:
            input_filepath: Путь к исходному файлу (относительно docs/)
            prompt: Промт для LLM модели
            llm_client: Клиент LLM (OpenAIClient)
        
        Returns:
            Путь к обработанному файлу
        """
        # Формируем полный путь к исходному файлу
        if not input_filepath.startswith('docs/'):
            input_filepath = f"docs/{input_filepath}"
        
        if not os.path.exists(input_filepath):
            raise FileNotFoundError(f"Файл не найден: {input_filepath}")
        
        print(f"📖 Обрабатываем документ: {input_filepath}")
        self.logger.info(f"Начинаем обработку документа: {input_filepath}")
        
        # Загружаем документ
        self.load_document(input_filepath)
        
        # Извлекаем текст из документа
        document_text = self._extract_text_from_document()
        self.logger.info(f"Извлечен текст из документа, длина: {len(document_text)} символов")
        
        # Отправляем в LLM
        if llm_client:
            print("🤖 Отправляем текст в LLM...")
            self.logger.info("Отправляем текст в LLM")
            modified_text = self._process_with_llm(document_text, prompt, llm_client)
            if modified_text:
                # Применяем изменения к документу
                self.logger.info("Применяем изменения от LLM к документу")
                self._apply_llm_changes(modified_text)
            else:
                print("❌ LLM не вернул результат")
                self.logger.error("LLM не вернул результат")
                return None
        else:
            print("❌ LLM клиент не предоставлен")
            self.logger.error("LLM клиент не предоставлен")
            return None
        
        # Генерируем имя выходного файла: "new_" + старое имя
        base_name = os.path.basename(input_filepath)
        output_filename = f"new_{base_name}"
        
        # Убеждаемся, что папка new_docs существует
        os.makedirs("new_docs", exist_ok=True)
        output_path = f"new_docs/{output_filename}"
        
        # Сохраняем обработанный документ
        self.save_document(output_path)
        
        print(f"✅ Обработка завершена: {output_path}")
        self.logger.info(f"Обработка документа завершена успешно: {output_path}")
        return output_path
    
    def _extract_text_from_document(self) -> str:
        """Извлекает весь текст из документа для отправки в LLM с уникальными маркерами"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        text_parts = []
        self.element_markers = []  # Сохраняем маркеры для последующего использования
        
        # Извлекаем текст из параграфов (включая пустые)
        for i, paragraph in enumerate(self.document.paragraphs):
            marker = f"PARA_{i:04d}"
            self.element_markers.append(('paragraph', i, marker))
            text_parts.append(f"[{marker}] {paragraph.text.strip()}")
        
        # Извлекаем текст из таблиц (по одной строке на ячейку)
        table_idx = 0
        for table in self.document.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    marker = f"CELL_{table_idx:02d}_{row_idx:02d}_{cell_idx:02d}"
                    self.element_markers.append(('table_cell', table_idx, row_idx, cell_idx, marker))
                    cell_text = cell.text.strip()
                    text_parts.append(f"[{marker}] {cell_text}")
            table_idx += 1
        
        self.logger.info(f"Извлечено элементов: {len(text_parts)} (параграфов: {len(self.document.paragraphs)}, таблиц: {len(self.document.tables)})")
        return "\n".join(text_parts)
    
    
    def _process_with_llm(self, document_text: str, prompt: str, llm_client) -> str:
        """Отправляет текст в LLM и получает измененный результат"""
        try:
            self.logger.info("Начинаем обработку текста в LLM")
            
            # Подсчитываем количество строк в оригинальном тексте
            original_lines = document_text.split('\n')
            line_count = len(original_lines)
            
            # Формируем полный промт
            full_prompt = f"""
{prompt}

Текст документа для обработки (всего {line_count} строк):
{document_text}

КРИТИЧЕСКИ ВАЖНО: 
1. Верни ТОЛЬКО измененный текст документа
2. Сохрани ТОЧНО {line_count} строк - ни больше, ни меньше
3. Каждая строка должна соответствовать оригинальной строке по порядку
4. Если строка пустая в оригинале, оставь её пустой
5. Не добавляй дополнительных комментариев, объяснений или форматирования
6. Не изменяй количество строк в документе
7. Верни результат в том же формате: каждая строка на новой строке
8. НЕ разбивай длинные строки на несколько строк
9. НЕ объединяй короткие строки в одну
10. Сохрани точную структуру документа
11. ОБЯЗАТЕЛЬНО сохрани все маркеры [PARA_XXXX] и [CELL_XX_XX_XX] в начале каждой строки
12. Изменяй только текст после маркера, НЕ трогай сам маркер
"""
            
            # Подготавливаем сообщения для LLM
            messages = llm_client.prepare_messages(full_prompt)
            self.logger.info(f"Отправляем запрос в LLM, длина промта: {len(full_prompt)} символов")
            
            # Отправляем запрос
            result = llm_client.generate(messages)
            
            if result and result[0]:
                modified_text = result[0]
                self.logger.info(f"LLM вернул результат, длина: {len(modified_text)} символов")
                
                # Логируем первые несколько строк для отладки
                lines = modified_text.split('\n')
                self.logger.info(f"LLM вернул {len(lines)} строк")
                for i, line in enumerate(lines[:5]):  # Показываем первые 5 строк
                    self.logger.info(f"Строка {i+1}: '{line}'")
                if len(lines) > 5:
                    self.logger.info(f"... и еще {len(lines) - 5} строк")
                
                return modified_text
            else:
                print("❌ LLM не вернул результат")
                self.logger.error("LLM не вернул результат")
                return None
                
        except Exception as e:
            print(f"❌ Ошибка при работе с LLM: {e}")
            self.logger.error(f"Ошибка при работе с LLM: {e}")
            return None
    
    def _apply_llm_changes(self, modified_text: str):
        """Применяет изменения от LLM к документу с сохранением форматирования"""
        if not self.document:
            raise ValueError("Документ не загружен!")

        print("🔧 Применяем изменения от LLM...")
        self.logger.info(f"Применяем изменения от LLM, длина текста: {len(modified_text)} символов")

        # Разбиваем измененный текст на строки, сохраняя пустые строки
        modified_lines = [line.rstrip() for line in modified_text.split('\n')]
        self.logger.info(f"LLM вернул {len(modified_lines)} строк")

        # Создаем словарь для быстрого поиска по маркерам
        marker_to_text = {}
        lines_without_markers = []
        
        for i, line in enumerate(modified_lines):
            if line.startswith('[') and ']' in line:
                marker_end = line.find(']')
                marker = line[1:marker_end]
                text_content = line[marker_end + 1:].strip()
                marker_to_text[marker] = text_content
            else:
                # Fallback для строк без маркеров
                lines_without_markers.append((i, line))
                self.logger.warning(f"Строка {i+1} без маркера: {line[:50]}...")

        # Валидация: проверяем соответствие количества элементов
        expected_elements = len(self.element_markers)
        actual_elements = len(modified_lines)
        
        if expected_elements != actual_elements:
            error_msg = f"Несоответствие количества элементов: ожидалось {expected_elements}, получено {actual_elements}"
            self.logger.error(error_msg)
            raise ValueError(error_msg)
        
        # Проверяем, есть ли строки без маркеров для fallback
        use_fallback = len(lines_without_markers) > 0
        if use_fallback:
            self.logger.warning(f"Используем fallback режим для {len(lines_without_markers)} строк без маркеров")

        paragraphs_updated = 0
        tables_updated = 0

        # Обновляем элементы по маркерам или позиционно (fallback)
        line_index = 0
        for element_info in self.element_markers:
            if element_info[0] == 'paragraph':
                _, para_idx, marker = element_info
                if marker in marker_to_text:
                    new_text = marker_to_text[marker]
                elif use_fallback and line_index < len(modified_lines):
                    new_text = modified_lines[line_index]
                    self.logger.info(f"Fallback: используем позиционный подход для параграфа {para_idx}")
                else:
                    self.logger.warning(f"Маркер {marker} не найден в ответе LLM")
                    continue
                
                paragraph = self.document.paragraphs[para_idx]
                self._update_paragraph_text(paragraph, new_text)
                paragraphs_updated += 1
            
            elif element_info[0] == 'table_cell':
                _, table_idx, row_idx, cell_idx, marker = element_info
                if marker in marker_to_text:
                    new_text = marker_to_text[marker]
                elif use_fallback and line_index < len(modified_lines):
                    new_text = modified_lines[line_index]
                    self.logger.info(f"Fallback: используем позиционный подход для ячейки {table_idx}-{row_idx}-{cell_idx}")
                else:
                    self.logger.warning(f"Маркер {marker} не найден в ответе LLM")
                    continue
                
                table = self.document.tables[table_idx]
                cell = table.rows[row_idx].cells[cell_idx]
                # Обновляем все параграфы в ячейке
                for paragraph in cell.paragraphs:
                    self._update_paragraph_text(paragraph, new_text)
                tables_updated += 1
            
            line_index += 1

        self.logger.info(f"Обновлено параграфов: {paragraphs_updated}, ячеек таблиц: {tables_updated}")
        print("✅ Изменения от LLM применены")
    
    def apply_changes(self, changes: Dict[str, str]) -> 'WordProcessor':
        """Применяет изменения к документу с сохранением форматирования"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        print("🔧 Применяем изменения...")
        changes_made = 0
        
        # Изменяем параграфы
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                original_text = paragraph.text
                new_text = original_text
                
                # Применяем замены
                for old_text, new_text_replacement in changes.items():
                    if old_text in new_text:
                        new_text = new_text.replace(old_text, new_text_replacement)
                        print(f"  🔄 Заменяем: '{old_text}' -> '{new_text_replacement}'")
                        changes_made += 1
                
                # Обновляем текст с сохранением форматирования
                if new_text != original_text:
                    self._update_paragraph_text(paragraph, new_text)
        
        # Изменяем таблицы
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            original_text = paragraph.text
                            new_text = original_text
                            
                            # Применяем замены
                            for old_text, new_text_replacement in changes.items():
                                if old_text in new_text:
                                    new_text = new_text.replace(old_text, new_text_replacement)
                                    print(f"  🔄 В таблице заменяем: '{old_text}' -> '{new_text_replacement}'")
                                    changes_made += 1
                            
                            # Обновляем текст
                            if new_text != original_text:
                                self._update_paragraph_text(paragraph, new_text)
        
        print(f"✅ Изменения применены: {changes_made} замен")
        return self
    
    def _update_paragraph_text(self, paragraph, new_text: str):
        """Обновляет текст параграфа с сохранением форматирования"""
        if paragraph.runs:
            # Сохраняем форматирование первого run
            first_run = paragraph.runs[0]
            first_run.text = new_text
            
            # Очищаем остальные runs
            for run in paragraph.runs[1:]:
                run.text = ''
        else:
            # Если нет runs, просто заменяем текст
            paragraph.text = new_text
    
    def save_document(self, output_path: str) -> 'WordProcessor':
        """Сохраняет документ"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        print(f"💾 Сохраняем документ: {output_path}")
        self.logger.info(f"Сохраняем документ: {output_path}")
        try:
            self.document.save(output_path)
            print("✅ Документ сохранен успешно")
            self.logger.info(f"Документ сохранен успешно: {output_path}")
        except Exception as e:
            self.logger.error(f"Ошибка при сохранении документа {output_path}: {e}")
            raise
        return self
    
    def show_document_info(self):
        """Показывает информацию о документе"""
        if not self.document:
            raise ValueError("Документ не загружен!")
        
        print("\n📋 Информация о документе:")
        print("-" * 40)
        
        # Подсчитываем элементы
        paragraphs = [p for p in self.document.paragraphs if p.text.strip()]
        tables = self.document.tables
        
        print(f"Параграфов: {len(paragraphs)}")
        print(f"Таблиц: {len(tables)}")
        
        # Показываем форматирование
        print("\n🎨 Форматирование:")
        for i, paragraph in enumerate(paragraphs[:5]):  # Показываем первые 5
            print(f"\nПараграф {i+1}:")
            print(f"  Текст: {paragraph.text[:50]}...")
            print(f"  Выравнивание: {paragraph.alignment}")
            
            if paragraph.runs:
                print("  Форматирование:")
                for j, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        print(f"    Run {j+1}: жирный={run.bold}, курсив={run.italic}")

def main():
    """Демонстрация работы процессора с LLM"""
    print("🔧 ПРОЦЕССОР WORD ДОКУМЕНТОВ С LLM")
    print("=" * 50)
    
    # Создаем процессор
    processor = WordProcessor()
    
    try:
        # Проверяем наличие файлов в папке docs
        docs_files = [f for f in os.listdir("docs") if f.endswith(('.docx', '.doc'))]
        
        if not docs_files:
            print("📁 Папка docs пуста. Поместите .docx файлы в папку docs/ для обработки.")
            print("💡 Пример использования:")
            print("   from word_processor import WordProcessor")
            print("   from llm import OpenRouterClient")
            print("   processor = WordProcessor()")
            print("   llm_client = OpenRouterClient()  # Использует API_KEY из env")
            print("   result = processor.process_document('мой_файл.docx', 'промт для LLM', llm_client)")
            return
        
        # Берем первый найденный файл для демонстрации
        demo_file = docs_files[0]
        print(f"📄 Найден файл для демонстрации: {demo_file}")
        
        # Пример промта для LLM
        prompt = """
Измени текст документа согласно следующим требованиям:
1. Замени все упоминания "философия" на "философствование"
2. Сохрани структуру документа и форматирование
Верни итоговый текст целиком.
"""
        
        print("💡 Для полной работы с LLM необходимо:")
        print("   1. Настроить API_KEY в переменных окружения")
        print("   2. Создать экземпляр OpenRouterClient")
        print("   3. Передать его в process_document()")
        print("\n📝 Пример промта:")
        print(prompt)
        
        # Демонстрация без LLM (только показываем структуру)
        print(f"\n🔧 Структура вызова:")
        print(f"   processor.process_document('{demo_file}', prompt, llm_client)")
        
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    main()
